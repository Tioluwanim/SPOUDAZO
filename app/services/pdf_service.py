"""
pdf_service.py - Handles PDF upload, storage, and retrieval.
Responsible for saving files to disk and managing document state.
"""

import uuid
import shutil
from pathlib import Path
from datetime import datetime

from app.config import (
    UPLOAD_DIR,
    PROCESSED_DIR,
    THUMBNAIL_DIR,
    R2_CACHE_DIR,
    STORAGE_PROVIDER,
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB,
    ALLOWED_EXTENSIONS,
)
from app.db.repository import repository
from app.models.schemas import (
    ProcessedDocument,
    DocumentMetadata,
    DocumentStatus,
    UploadResponse,
    ErrorResponse,
)
from app.services.storage import storage_service
from app.utils.logger import get_logger, ServiceLogger

logger = get_logger(__name__)


# ── PDF Service Class ─────────────────────────────────────────────────────────

class PDFService:
    """
    Handles all file-level operations:
    - Validating uploaded PDFs
    - Saving to disk
    - Loading processed document state
    - Listing and deleting documents
    """

    def __init__(self):
        self.upload_dir    = UPLOAD_DIR
        self.processed_dir = PROCESSED_DIR
        self.repository    = repository
        logger.info("PDFService initialised")

    # ── Upload ────────────────────────────────────────────────────────────────

    def save_upload(
        self,
        file_bytes: bytes,
        filename: str,
        checksum: str | None = None,
    ) -> tuple[ProcessedDocument, None] | tuple[None, ErrorResponse]:
        """
        Validates and saves an uploaded file to disk.
        Supports: PDF, DOCX, DOC, TXT, XLSX, XLS, CSV.
        Non-PDF files are converted to a text-based PDF wrapper so the
        rest of the pipeline (extraction → embedding → RAG) works unchanged.

        `checksum` (sha256 of the original upload bytes, computed by the
        caller before any PDF conversion) is stored on the Document record
        so a later upload of the same content can be recognised via
        repository.get_ready_document_by_checksum() instead of re-running
        extraction/embedding/indexing on identical material.

        Returns:
            (ProcessedDocument, None) on success.
            (None, ErrorResponse)     on validation failure.
        """
        slog = ServiceLogger("pdf_service")

        suffix = Path(filename).suffix.lower()
        if suffix not in ALLOWED_EXTENSIONS:
            msg = (
                f"Unsupported file type '{suffix}'. "
                f"Supported: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
            )
            slog.warning(msg)
            return None, ErrorResponse(error="Invalid file type", detail=msg)

        file_size = len(file_bytes)
        if file_size == 0:
            return None, ErrorResponse(error="Empty file", detail="Uploaded file is empty.")
        if file_size > MAX_FILE_SIZE_BYTES:
            msg = (
                f"File size {file_size / (1024*1024):.1f} MB exceeds "
                f"the {MAX_FILE_SIZE_MB} MB limit."
            )
            return None, ErrorResponse(error="File too large", detail=msg)

        # ── For non-PDF files: extract text and wrap as a .pdf on disk ────────
        if suffix != ".pdf":
            try:
                text = _extract_text_from_file(file_bytes, suffix, filename)
                if not text.strip():
                    return None, ErrorResponse(
                        error="Empty content",
                        detail=f"Could not extract any text from '{filename}'.",
                    )
                # Wrap as minimal PDF using PyMuPDF
                import fitz
                pdf_doc  = fitz.open()
                page     = pdf_doc.new_page()
                # Insert text with automatic wrapping
                page.insert_textbox(
                    fitz.Rect(50, 50, 562, 792),
                    text,
                    fontsize  = 9,
                    fontname  = "helv",
                    align     = 0,
                )
                # If text overflows one page, add more pages
                if len(text) > 4000:
                    chunks = [text[i:i+4000] for i in range(4000, len(text), 4000)]
                    for chunk in chunks:
                        pg = pdf_doc.new_page()
                        pg.insert_textbox(
                            fitz.Rect(50, 50, 562, 792),
                            chunk, fontsize=9, fontname="helv", align=0,
                        )
                file_bytes = pdf_doc.tobytes()
                pdf_doc.close()
                # Keep original name but save as .pdf
                filename   = Path(filename).stem + ".pdf"
                slog.info(
                    "Converted '%s' (%s) → PDF (%d bytes, %d chars text)",
                    filename, suffix, len(file_bytes), len(text),
                )
            except Exception as e:
                return None, ErrorResponse(
                    error="Conversion failed",
                    detail=f"Could not convert '{filename}' to PDF: {e}",
                )

        # ── Validate PDF magic bytes ──────────────────────────────────────────
        if not file_bytes.startswith(b"%PDF"):
            return None, ErrorResponse(
                error="Invalid PDF",
                detail="File does not appear to be a valid PDF.",
            )

        # ── Save via StorageService ──────────────────────────────────────────
        doc_id     = str(uuid.uuid4())
        safe_name  = self._sanitize_filename(filename)
        storage_key = f"{doc_id}_{safe_name}"

        slog = ServiceLogger("pdf_service", doc_id=doc_id)

        try:
            storage_url = storage_service.upload(storage_key, file_bytes, content_type="application/pdf")
        except Exception as e:
            msg = f"Failed to save file: {e}"
            slog.error(msg)
            return None, ErrorResponse(error="Storage error", detail=msg, doc_id=doc_id)

        # local_path stays a real, directly-readable filesystem path for the
        # local provider (StorageService writes to UPLOAD_DIR/storage_key,
        # exactly where this used to write it directly) - every existing
        # consumer of local_path/file_path (extraction, thumbnails) keeps
        # working unchanged. For the r2 provider it's not a real local path;
        # get_local_file_path() below is what those consumers should move to.
        dest_path = self.upload_dir / storage_key
        slog.info(f"Saved '{filename}' via {STORAGE_PROVIDER} provider → key={storage_key} ({file_size:,} bytes)")

        # ── Build ProcessedDocument ───────────────────────────────────────────
        doc = ProcessedDocument(
            doc_id    = doc_id,
            filename  = filename,
            file_path = str(dest_path),
            status    = DocumentStatus.UPLOADED,
            metadata  = DocumentMetadata(file_size_bytes=file_size),
        )

        self.repository.create_document(
            doc_id         = doc_id,
            filename       = filename,
            local_path     = str(dest_path),
            file_size_bytes= file_size,
            mime_type      = "application/pdf",
            source_folder  = None,
            drive_file_id  = None,
            checksum       = checksum,
            storage_provider = STORAGE_PROVIDER,
            storage_key      = storage_key,
            storage_url      = storage_url,
            modified_time  = None,
            source         = "upload",
        )

        slog.info(f"Document record created — status={doc.status.value}")
        return doc, None

    def convert_to_pdf_bytes(
        self,
        file_bytes: bytes,
        filename: str,
    ) -> tuple[bytes, str]:
        """
        Convert a supported non-PDF document to PDF bytes.

        Returns:
            A tuple of (pdf_bytes, output_filename).
        """
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return file_bytes, filename

        text = _extract_text_from_file(file_bytes, suffix, filename)
        if not text.strip():
            raise RuntimeError(f"Could not extract text from '{filename}'.")

        import fitz

        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_textbox(
            fitz.Rect(50, 50, 562, 792),
            text,
            fontsize=9,
            fontname="helv",
            align=0,
        )
        if len(text) > 4000:
            chunks = [text[i:i+4000] for i in range(4000, len(text), 4000)]
            for chunk in chunks:
                pg = pdf_doc.new_page()
                pg.insert_textbox(
                    fitz.Rect(50, 50, 562, 792),
                    chunk,
                    fontsize=9,
                    fontname="helv",
                    align=0,
                )
        pdf_bytes = pdf_doc.tobytes()
        pdf_doc.close()
        return pdf_bytes, Path(filename).stem + ".pdf"

    # ── Load / Save State ─────────────────────────────────────────────────────

    def load_document(self, doc_id: str) -> ProcessedDocument | None:
        """
        Loads a ProcessedDocument from its JSON state file.

        Args:
            doc_id: The document UUID.

        Returns:
            ProcessedDocument if found, None otherwise.
        """
        doc = self.repository.load_processed_document(doc_id)
        if doc:
            logger.debug(f"[{doc_id}] Loaded document from DB — status={doc.status.value}")
            return doc

        state_path = self._state_path(doc_id)
        if not state_path.exists():
            logger.warning(f"[{doc_id}] State file not found: {state_path}")
            return None

        try:
            doc = ProcessedDocument.model_validate_json(state_path.read_text(encoding="utf-8"))
            logger.debug(f"[{doc_id}] Loaded document from JSON state — status={doc.status.value}")
            return doc
        except Exception as e:
            logger.error(f"[{doc_id}] Failed to load state: {e}", exc_info=True)
            return None

    def save_document(self, doc: ProcessedDocument) -> bool:
        """
        Persists a ProcessedDocument to the database.
        Updates updated_at timestamp automatically.

        Returns:
            True on success, False on failure.
        """
        doc.updated_at = datetime.utcnow()
        return self.repository.update_document(doc)

    def update_status(
        self,
        doc_id: str,
        status: DocumentStatus,
        error_message: str | None = None,
    ) -> bool:
        """
        Convenience method to update only the status of a document.

        Args:
            doc_id:        Document UUID.
            status:        New DocumentStatus value.
            error_message: Optional error detail (set on FAILED status).

        Returns:
            True on success, False if document not found.
        """
        doc = self.load_document(doc_id)
        if not doc:
            return False

        doc.status = status
        if error_message:
            doc.error_message = error_message

        logger.info(f"[{doc_id}] Status → {status.value}")
        return self.save_document(doc)

    # ── List Documents ────────────────────────────────────────────────────────

    def list_documents(self) -> list[dict]:
        """
        Returns a list of all document summaries from the database.
        Used to populate the sidebar in the Streamlit UI.
        """
        try:
            return self.repository.list_documents()
        except Exception as e:
            logger.error("Failed to list documents: %s", e, exc_info=True)
            return []

    # ── Delete Document ───────────────────────────────────────────────────────

    def delete_document(self, doc_id: str) -> bool:
        """
        Deletes all files and database state for a document.

        Args:
            doc_id: Document UUID.

        Returns:
            True if anything was deleted, False if nothing found.
        """
        slog   = ServiceLogger("pdf_service", doc_id=doc_id)
        doc    = self.load_document(doc_id)
        found  = False

        db_doc = self.repository.get_document_by_doc_id(doc_id)
        if db_doc and db_doc.storage_key:
            # Route through StorageService rather than unlinking
            # doc.file_path directly - for an r2-provider document,
            # file_path is either a local scratch cache (safe to also
            # remove below) or nothing at all; the actual object only
            # gets deleted by going through the provider that owns it.
            try:
                if storage_service.delete(db_doc.storage_key):
                    slog.info(f"Deleted original file from {db_doc.storage_provider or 'local'} storage")
                    found = True
            except Exception as e:
                slog.error(f"Failed to delete from storage: {e}")
        elif doc and Path(doc.file_path).exists():
            # Legacy row with no storage_key (predates StorageService) -
            # same direct unlink this always did.
            Path(doc.file_path).unlink()
            slog.info("Deleted uploaded PDF (legacy direct path)")
            found = True

        # Clean up the local R2 scratch cache too, if get_local_file_path
        # ever populated one for this document.
        cache_path = R2_CACHE_DIR / f"{doc_id}_{Path(doc.file_path).name}" if doc else None
        if cache_path and cache_path.exists():
            cache_path.unlink()
            found = True

        # rag_service.delete_index (not a raw shutil.rmtree here) so its
        # in-memory FAISS/BM25 caches are invalidated too - deleting only
        # the files on disk would leave chat/search still "finding" this
        # document's content out of a stale cache until the next restart.
        from app.services.rag_service import rag_service
        if rag_service.delete_index(doc_id):
            found = True

        if self.repository.delete_document(doc_id):
            slog.info("Deleted document record from database")
            found = True

        state_path = self._state_path(doc_id)
        if state_path.exists():
            state_path.unlink()
            slog.info("Deleted legacy state file")
            found = True

        return found

    # ── Helpers ───────────────────────────────────────────────────────────────

    def get_upload_response(self, doc: ProcessedDocument) -> UploadResponse:
        """Builds an UploadResponse from a ProcessedDocument."""
        return UploadResponse(
            doc_id    = doc.doc_id,
            filename  = doc.filename,
            file_size = doc.metadata.file_size_bytes,
            status    = doc.status,
        )

    def document_exists(self, doc_id: str) -> bool:
        """Returns True if the document exists in the database or legacy state file."""
        return self.repository.get_document_by_doc_id(doc_id) is not None or self._state_path(doc_id).exists()

    def is_ready(self, doc_id: str) -> bool:
        """Returns True if document is fully processed and ready for chat."""
        doc = self.load_document(doc_id)
        return doc is not None and doc.status == DocumentStatus.READY

    def get_local_file_path(self, doc_id: str) -> Path | None:
        """
        Returns a real, directly-readable local filesystem path for this
        document's original file - even when storage_provider="r2".

        For local-provider documents (including every document created
        before this feature existed, since storage_provider is NULL on
        those rows and NULL is treated as "local"), this is just
        Document.local_path - no network call, no behavior change.

        For r2-provider documents, R2 is the durable source of truth but
        the existing extraction/thumbnail code is written against real
        file paths (fitz.open(path), etc.) - rewriting that code to
        stream everything is a much bigger change than this app's
        existing pipeline needs right now. Instead, this downloads once
        into R2_CACHE_DIR (ephemeral scratch space, safe to lose - it's
        regenerated from R2 on demand) and returns that cached path on
        every subsequent call.
        """
        document = self.repository.get_document_by_doc_id(doc_id)
        if document is None:
            return None

        if document.storage_provider in (None, "", "local"):
            path = Path(document.local_path)
            return path if path.exists() else None

        # r2 (or any future non-local provider)
        cache_path = R2_CACHE_DIR / f"{doc_id}_{Path(document.local_path).name}"
        if cache_path.exists():
            return cache_path

        if not document.storage_key:
            logger.error(f"[{doc_id}] storage_provider={document.storage_provider!r} but no storage_key set")
            return None

        try:
            data = storage_service.download(document.storage_key)
        except Exception as e:
            logger.error(f"[{doc_id}] Failed to download from {document.storage_provider}: {e}")
            return None

        cache_path.parent.mkdir(parents=True, exist_ok=True)
        cache_path.write_bytes(data)
        return cache_path

    def get_page_thumbnail(self, doc_id: str, page_number: int, target_width: int = 320) -> Path | None:
        """
        Renders (and caches to disk) a low-res PNG of one page, for the
        reader's page-thumbnail strip. Generated on first request, not
        during upload processing - upload should stay fast, and most
        pages of most documents are never actually opened via the
        thumbnail strip, so rendering all of them upfront would mostly
        be wasted work.

        Returns None if the document/page doesn't exist or rendering
        fails (e.g. a non-PDF-derived document with no real pages) -
        callers treat that as "no thumbnail available", not an error.
        """
        cache_path = THUMBNAIL_DIR / doc_id / f"{page_number}.png"
        if cache_path.exists():
            return cache_path

        file_path = self.get_local_file_path(doc_id)
        if file_path is None:
            return None

        try:
            from app.services.extraction_service import _get_fitz_module
            fitz = _get_fitz_module()
        except ImportError:
            logger.warning(f"[{doc_id}] PyMuPDF unavailable — cannot render thumbnails")
            return None

        try:
            with fitz.open(str(file_path)) as pdf:
                if not (1 <= page_number <= pdf.page_count):
                    return None
                page = pdf[page_number - 1]  # fitz pages are 0-indexed
                zoom = target_width / page.rect.width if page.rect.width else 1.0
                pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                cache_path.parent.mkdir(parents=True, exist_ok=True)
                pixmap.save(str(cache_path))
                return cache_path
        except Exception as e:
            logger.warning(f"[{doc_id}] Thumbnail render failed for page {page_number}: {e}")
            return None

    # ── Private ───────────────────────────────────────────────────────────────

    def _state_path(self, doc_id: str) -> Path:
        """Returns the path to the JSON state file for a document."""
        return self.processed_dir / f"{doc_id}.json"

    def _save_document_state(self, doc: ProcessedDocument) -> bool:
        """Writes the ProcessedDocument as JSON to processed_dir."""
        state_path = self._state_path(doc.doc_id)
        try:
            state_path.write_text(
                doc.model_dump_json(indent=2),
                encoding="utf-8",
            )
            return True
        except OSError as e:
            logger.error(f"[{doc.doc_id}] Failed to save state: {e}", exc_info=True)
            return False

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """
        Removes unsafe characters from a filename.
        Keeps alphanumerics, dots, hyphens and underscores.
        """
        safe = "".join(
            c if c.isalnum() or c in (".", "-", "_") else "_"
            for c in filename
        )
        # Collapse multiple underscores
        while "__" in safe:
            safe = safe.replace("__", "_")
        return safe.strip("_") or "document.pdf"


# ── Module-level singleton ────────────────────────────────────────────────────
# Services import this instance directly — no need to instantiate.
pdf_service = PDFService()


# ── File-to-text converters ───────────────────────────────────────────────────

def _extract_text_from_file(file_bytes: bytes, suffix: str, filename: str) -> str:
    """
    Extract plain text from a non-PDF file.
    Supports: .docx, .doc, .txt, .xlsx, .xls, .csv

    Returns extracted text as a single string.
    Raises on unrecoverable errors.
    """
    import io as _io

    # ── Plain text ────────────────────────────────────────────────────────────
    if suffix == ".txt":
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(enc)
            except UnicodeDecodeError:
                continue
        return file_bytes.decode("utf-8", errors="replace")

    # ── CSV ───────────────────────────────────────────────────────────────────
    if suffix == ".csv":
        import csv as _csv
        text_io = _io.StringIO(file_bytes.decode("utf-8-sig", errors="replace"))
        reader  = _csv.reader(text_io)
        rows    = list(reader)
        if not rows:
            return ""
        # Build readable text: header as labels, rows as key:value lines
        header = rows[0]
        lines  = ["\t".join(header)]
        for row in rows[1:]:
            pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip()]
            lines.append("  |  ".join(pairs))
        return "\n".join(lines)

    # ── DOCX ──────────────────────────────────────────────────────────────────
    if suffix in (".docx",):
        try:
            from docx import Document as _DocxDoc
            docx    = _DocxDoc(_io.BytesIO(file_bytes))
            parts   = [p.text for p in docx.paragraphs if p.text.strip()]
            # Also extract tables
            for table in docx.tables:
                for row in table.rows:
                    row_text = "  |  ".join(
                        c.text.strip() for c in row.cells if c.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n\n".join(parts)
        except ImportError:
            raise RuntimeError("python-docx not installed. Add it to requirements.txt.")

    # ── DOC (legacy Word) ─────────────────────────────────────────────────────
    if suffix == ".doc":
        # Try antiword-style extraction via python-docx (sometimes works)
        # Otherwise fall back to raw text extraction
        try:
            from docx import Document as _DocxDoc
            docx  = _DocxDoc(_io.BytesIO(file_bytes))
            parts = [p.text for p in docx.paragraphs if p.text.strip()]
            return "\n\n".join(parts)
        except Exception:
            # Raw fallback: extract printable ASCII from binary
            raw = file_bytes.decode("latin-1", errors="replace")
            return "\n".join(
                line for line in raw.splitlines()
                if len(line.strip()) > 20 and line.strip().isprintable()
            )

    # ── XLSX / XLS ────────────────────────────────────────────────────────────
    if suffix in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb    = openpyxl.load_workbook(_io.BytesIO(file_bytes), data_only=True)
            parts : list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"=== Sheet: {sheet_name} ===")
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(v) if v is not None else "" for v in row]
                    non_empty = [v for v in row_vals if v.strip()]
                    if non_empty:
                        parts.append("  |  ".join(non_empty))
            return "\n".join(parts)
        except ImportError:
            raise RuntimeError("openpyxl not installed. Add it to requirements.txt.")
